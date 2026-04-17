from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import Workbook
from pypdf import PdfWriter

import markbridge.routing.runtime as runtime_module
from markbridge.parsers.base import ParseRequest
from markbridge.pipeline import PipelineRequest, run_pipeline
from markbridge.routing.runtime import executable_candidates_for_format
from markbridge.parsers.basic import _blocks_from_markdown
from markbridge.shared.ir import DocumentFormat


def test_docx_pipeline_smoke() -> None:
    doc = DocxDocument()
    doc.add_heading("Sample Heading", level=1)
    doc.add_paragraph("Hello world")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert result.parser_id == "python-docx"
    assert result.decision.value == "accept"
    assert "Sample Heading" in result.metadata["markdown"]
    assert len(result.trace.events) >= 20
    assert result.artifacts.markdown is not None
    assert Path(result.metadata["export_dir"]).exists()


def test_docx_pipeline_promotes_numbered_plain_paragraph_to_heading() -> None:
    doc = DocxDocument()
    doc.add_paragraph("1. 보장내용")
    doc.add_paragraph("이 문단은 보장 범위에 대한 설명을 담고 있습니다.")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert "## 1. 보장내용" in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "heading"
    assert first_block.metadata["chunk_boundary_reason"] == "heading_pattern.numbered"
    assert first_block.metadata["level"] == 2


def test_docx_pipeline_uses_deeper_heading_level_for_decimal_numbering() -> None:
    doc = DocxDocument()
    doc.add_paragraph("1.1 가입대상")
    doc.add_paragraph("세부 설명")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert "### 1.1 가입대상" in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.metadata["level"] == 3


def test_docx_pipeline_promotes_short_plain_title_to_heading_when_followed_by_body() -> None:
    doc = DocxDocument()
    doc.add_paragraph("보장내용")
    doc.add_paragraph("이 문단은 일반 본문이며 제목 바로 아래의 설명 역할을 합니다.")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert "## 보장내용" in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "heading"
    assert first_block.metadata["chunk_boundary_reason"] == "heading_pattern.short_title"


def test_docx_pipeline_does_not_promote_long_plain_body_sentence_to_heading() -> None:
    doc = DocxDocument()
    body = "이 문장은 일반 본문으로 작성되었고 제목이 아니라 설명 문단이기 때문에 heading으로 승격되면 안 됩니다."
    doc.add_paragraph(body)
    doc.add_paragraph("뒤 문단도 일반 설명입니다.")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert f"## {body}" not in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "paragraph"


def test_docx_pipeline_promotes_korean_section_marker_to_heading() -> None:
    doc = DocxDocument()
    doc.add_paragraph("제1장 보장내용")
    doc.add_paragraph("이 장에서는 보장 범위를 설명합니다.")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert "## 제1장 보장내용" in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "heading"
    assert first_block.metadata["chunk_boundary_reason"] == "heading_pattern.korean_section"


def test_docx_pipeline_promotes_parenthesized_number_line_to_heading() -> None:
    doc = DocxDocument()
    doc.add_paragraph("1) 우편접수")
    doc.add_paragraph("이 문단은 세부 안내를 설명합니다.")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert "## 1) 우편접수" in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "heading"
    assert first_block.metadata["chunk_boundary_reason"] == "heading_pattern.numbered"
    assert first_block.metadata["level"] == 2


def test_docx_pipeline_promotes_selected_closed_number_section_heading() -> None:
    doc = DocxDocument()
    doc.add_paragraph("3) 대리인")
    doc.add_paragraph("대리인 기준과 구비서류를 아래에서 설명합니다.")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert "## 3) 대리인" in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "heading"
    assert first_block.metadata["chunk_boundary_reason"] == "heading_pattern.numbered"


def test_docx_pipeline_promotes_circled_number_section_heading_when_followed_by_table() -> None:
    doc = DocxDocument()
    doc.add_paragraph("① 보험계약조회 및 보험료납입")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "구분"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "업무"
    table.cell(1, 1).text = "조회 및 납입"
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert "## ① 보험계약조회 및 보험료납입" in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "heading"
    assert first_block.metadata["chunk_boundary_reason"] == "heading_pattern.circled_number_section"


def test_docx_pipeline_does_not_promote_circled_number_body_line_without_section_context() -> None:
    doc = DocxDocument()
    doc.add_paragraph("① 계약정보 : 계약일, 만기일, 납입상태")
    doc.add_paragraph("뒤 문단은 일반 설명 본문입니다.")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert "## ① 계약정보 : 계약일, 만기일, 납입상태" not in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "paragraph"


def test_docx_pipeline_promotes_circled_number_section_sequence_without_table() -> None:
    doc = DocxDocument()
    doc.add_paragraph("① 보험계약조회 및 보험료납입")
    doc.add_paragraph("안내 본문")
    doc.add_paragraph("② 보험금신청")
    doc.add_paragraph("추가 안내 본문")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    markdown = result.metadata["markdown"]
    assert "## ① 보험계약조회 및 보험료납입" in markdown
    assert "## ② 보험금신청" in markdown


def test_docx_pipeline_promotes_custom_korean_title_style_to_heading() -> None:
    doc = DocxDocument()
    style = doc.styles.add_style("첫제목", WD_STYLE_TYPE.PARAGRAPH)
    paragraph = doc.add_paragraph("보험종목의 명칭")
    paragraph.style = style
    doc.add_paragraph("1형 : 무배당 종신보험 표준형")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    assert "## 보험종목의 명칭" in result.metadata["markdown"]
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "heading"
    assert first_block.metadata["chunk_boundary_reason"] == "heading_style.custom_title"
    assert first_block.metadata["level"] == 2


def test_docx_pipeline_maps_second_custom_title_style_to_deeper_heading() -> None:
    doc = DocxDocument()
    first_style = doc.styles.add_style("첫제목", WD_STYLE_TYPE.PARAGRAPH)
    second_style = doc.styles.add_style("두번째제목", WD_STYLE_TYPE.PARAGRAPH)
    p1 = doc.add_paragraph("보험종목의 명칭")
    p1.style = first_style
    p2 = doc.add_paragraph("1형 : 무배당 종신보험 표준형")
    p2.style = second_style
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    markdown = result.metadata["markdown"]
    assert "## 보험종목의 명칭" in markdown
    assert "### 1형 : 무배당 종신보험 표준형" in markdown


def test_xlsx_pipeline_smoke() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Stats"
    sheet["A1"] = "Header"
    sheet["A2"] = 100
    with NamedTemporaryFile(suffix=".xlsx", delete=False) as handle:
        path = Path(handle.name)
    workbook.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.XLSX))

    assert result.parser_id == "openpyxl"
    assert result.decision.value in {"accept", "degraded_accept"}
    assert "## Stats" in result.metadata["markdown"]
    assert "Header" in result.metadata["markdown"]
    assert result.artifacts.trace is not None


def test_xlsx_pipeline_emits_sheet_heading_block() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Claims Summary"
    sheet["A1"] = "Metric"
    sheet["B1"] = "Value"
    sheet["A2"] = "Loss Ratio"
    sheet["B2"] = 88
    with NamedTemporaryFile(suffix=".xlsx", delete=False) as handle:
        path = Path(handle.name)
    workbook.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.XLSX))

    assert result.metadata["markdown"].startswith("## Claims Summary")
    first_block = result.document.blocks[0]
    assert first_block.kind.value == "heading"
    assert first_block.metadata["chunk_boundary_reason"] == "sheet_name"


def test_xlsx_pipeline_emits_heading_for_each_sheet() -> None:
    workbook = Workbook()
    first = workbook.active
    first.title = "Summary"
    first["A1"] = "Header"
    first["A2"] = "Value"
    second = workbook.create_sheet("Details")
    second["A1"] = "Metric"
    second["A2"] = "Count"
    with NamedTemporaryFile(suffix=".xlsx", delete=False) as handle:
        path = Path(handle.name)
    workbook.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.XLSX))

    markdown = result.metadata["markdown"]
    assert "## Summary" in markdown
    assert "## Details" in markdown


def test_markdown_heading_blocks_carry_chunk_boundary_metadata() -> None:
    blocks = _blocks_from_markdown("## 보장내용\n\n본문", default_page_number=1)

    heading = blocks[0]
    assert heading.kind.value == "heading"
    assert heading.metadata["chunk_boundary_candidate"] is True
    assert heading.metadata["chunk_boundary_reason"] == "markdown_heading"
    assert heading.metadata["chunk_boundary_materialized"] is True


def test_docx_single_column_layout_table_is_flattened_into_chunkable_blocks() -> None:
    doc = DocxDocument()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = (
        "(1) 변경시 지급보류 계약은 확인 필요\n"
        "■ 계약자 = 채무자 -> 변경 불가\n"
        "(2) 보험계약대출이 있는 경우 확인 필요\n"
        "추가 안내 문구"
    )
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    markdown = result.metadata["markdown"]
    assert "> (1) 변경시 지급보류 계약은 확인 필요" in markdown
    assert "> ■ 계약자 = 채무자 -> 변경 불가" in markdown
    assert "> (2) 보험계약대출이 있는 경우 확인 필요" in markdown
    assert "| --- |" not in markdown
    note_block = result.document.blocks[0]
    assert note_block.kind.value == "note"
    assert note_block.metadata["box_preserved"] is True


def test_docx_multi_column_table_still_renders_as_markdown_table() -> None:
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "구분"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "우편접수"
    table.cell(1, 1).text = "가능"
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    markdown = result.metadata["markdown"]
    assert "| 구분 | 내용 |" in markdown
    assert "| 우편접수 | 가능 |" in markdown


def test_docx_table_carry_forward_fills_blank_parent_category_cells() -> None:
    doc = DocxDocument()
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "접수방법"
    table.cell(0, 1).text = "구분"
    table.cell(0, 2).text = "구비서류 안내"
    table.cell(1, 0).text = "내방"
    table.cell(1, 1).text = "모두 방문"
    table.cell(1, 2).text = "신분증"
    table.cell(2, 0).text = ""
    table.cell(2, 1).text = "대리인"
    table.cell(2, 2).text = "위임장"
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    markdown = result.metadata["markdown"]
    assert "| 내방 | 대리인 | 위임장 |" in markdown
    table_block = next(block for block in result.document.blocks if block.kind.value == "table")
    assert table_block.metadata["docx_carry_forward"] is True


def test_docx_table_collapses_horizontal_merge_duplicates() -> None:
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=3)
    merged_header = table.cell(0, 0).merge(table.cell(0, 1))
    merged_header.text = "접수방법"
    table.cell(0, 2).text = "구비서류 안내"
    merged_value = table.cell(1, 0).merge(table.cell(1, 1))
    merged_value.text = "우편/\nFC방문 접수"
    table.cell(1, 2).text = "신청서"
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    markdown = result.metadata["markdown"]
    assert "| 접수방법 | 구비서류 안내 |" in markdown
    assert "| 접수방법 | 접수방법 | 구비서류 안내 |" not in markdown
    assert "| 우편/\nFC방문 접수 | 신청서 |" in markdown


def test_docx_pipeline_preserves_paragraph_table_order() -> None:
    doc = DocxDocument()
    doc.add_paragraph("3. 대리인")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "구분"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "대리인 자격"
    table.cell(1, 1).text = "직계가족"
    doc.add_paragraph("4. 중요 확인 사항")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOCX))

    markdown = result.metadata["markdown"]
    assert markdown.index("## 3. 대리인") < markdown.index("| 구분 | 내용 |")
    assert markdown.index("| 구분 | 내용 |") < markdown.index("## 4. 중요 확인 사항")


def test_pdf_pipeline_smoke() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    with NamedTemporaryFile(suffix=".pdf", delete=False) as handle:
        path = Path(handle.name)
    with path.open("wb") as output:
        writer.write(output)

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.PDF))

    assert result.parser_id in {"docling", "pypdf"}
    assert result.status.value in {"degraded", "failed"}
    assert result.artifacts.metadata is not None


def test_doc_pipeline_is_held_when_no_route_exists() -> None:
    with NamedTemporaryFile(suffix=".doc", delete=False) as handle:
        path = Path(handle.name)
        handle.write(b"legacy-doc-placeholder")

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOC))

    assert result.parser_id is None
    assert result.decision.value == "hold"
    assert result.validation.issues


def test_executable_candidates_reflect_current_environment() -> None:
    assert executable_candidates_for_format(DocumentFormat.DOCX) == ("python-docx",)
    assert executable_candidates_for_format(DocumentFormat.XLSX) == ("openpyxl",)
    assert executable_candidates_for_format(DocumentFormat.PDF) == ("docling", "pypdf")


def test_doc_candidate_appears_when_libreoffice_enabled(monkeypatch) -> None:
    original = runtime_module.get_runtime_statuses

    def fake_statuses():
        statuses = original()
        statuses["libreoffice"] = runtime_module.RuntimeParserStatus("libreoffice", True, True)
        return statuses

    monkeypatch.setattr(runtime_module, "get_runtime_statuses", fake_statuses)
    assert executable_candidates_for_format(DocumentFormat.DOC) == ("libreoffice",)


def test_doc_candidate_appears_with_antiword_fallback_when_enabled(monkeypatch) -> None:
    original = runtime_module.get_runtime_statuses

    def fake_statuses():
        statuses = original()
        statuses["libreoffice"] = runtime_module.RuntimeParserStatus("libreoffice", False, False, "not installed")
        statuses["antiword"] = runtime_module.RuntimeParserStatus("antiword", True, True)
        return statuses

    monkeypatch.setattr(runtime_module, "get_runtime_statuses", fake_statuses)
    assert executable_candidates_for_format(DocumentFormat.DOC) == ("antiword",)


def test_hwp_candidate_appears_when_hwp5txt_enabled(monkeypatch) -> None:
    original = runtime_module.get_runtime_statuses

    def fake_statuses():
        statuses = original()
        statuses["hwp5txt"] = runtime_module.RuntimeParserStatus("hwp5txt", True, True)
        return statuses

    monkeypatch.setattr(runtime_module, "get_runtime_statuses", fake_statuses)
    assert executable_candidates_for_format(DocumentFormat.HWP) == ("hwp5txt",)


def test_docx_pipeline_ignores_disabled_markitdown_override() -> None:
    doc = DocxDocument()
    doc.add_heading("Override Heading", level=1)
    doc.add_paragraph("Override body")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    result = run_pipeline(
        PipelineRequest(
            source_path=path,
            document_format=DocumentFormat.DOCX,
            options={"parser_override": "markitdown", "llm_route_used": True},
        )
    )

    assert result.parser_id == "python-docx"
    assert result.decision.value == "accept"
    assert "Override Heading" in result.metadata["markdown"]


def test_hwp_pipeline_is_held_as_unimplemented() -> None:
    with NamedTemporaryFile(suffix=".hwp", delete=False) as handle:
        path = Path(handle.name)
        handle.write(b"hwp-placeholder")

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.HWP))

    assert result.parser_id is None
    assert result.decision.value == "hold"


def test_doc_pipeline_degrades_handoff_when_antiword_route_is_used(monkeypatch) -> None:
    import markbridge.inspection.basic as inspection_module
    import markbridge.parsers.basic as parser_module
    import markbridge.routing.runtime as runtime_module

    original_statuses = runtime_module.get_runtime_statuses

    def fake_statuses():
        statuses = original_statuses()
        statuses["libreoffice"] = runtime_module.RuntimeParserStatus("libreoffice", False, False, "not installed")
        statuses["antiword"] = runtime_module.RuntimeParserStatus(
            "antiword",
            True,
            True,
            None,
            supported_formats=(DocumentFormat.DOC,),
            route_kind="degraded_fallback",
        )
        return statuses

    monkeypatch.setattr(runtime_module, "get_runtime_statuses", fake_statuses)
    monkeypatch.setattr(inspection_module, "antiword_available", lambda: True)
    monkeypatch.setattr(
        parser_module,
        "extract_doc_text_with_antiword",
        lambda _path: parser_module.TextExtractionResult(
            succeeded=True,
            text="1. 안내사항\n\n대리인 접수 가능",
            message="DOC extracted with antiword text fallback.",
        ),
    )

    with NamedTemporaryFile(suffix=".doc", delete=False) as handle:
        path = Path(handle.name)
        handle.write(b"legacy-doc-placeholder")

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.DOC))

    assert result.parser_id == "antiword"
    assert result.decision.value == "degraded_accept"
    assert "degraded_parser_route" in result.handoff.reasons
    assert result.handoff.metadata["parser_route_kind"] == "degraded_fallback"


def test_hwp_pipeline_degrades_handoff_when_text_route_is_used(monkeypatch) -> None:
    import markbridge.inspection.basic as inspection_module
    import markbridge.parsers.basic as parser_module
    import markbridge.routing.runtime as runtime_module

    original_statuses = runtime_module.get_runtime_statuses

    def fake_statuses():
        statuses = original_statuses()
        statuses["hwp5txt"] = runtime_module.RuntimeParserStatus(
            "hwp5txt",
            True,
            True,
            None,
            supported_formats=(DocumentFormat.HWP,),
            route_kind="text_route",
        )
        return statuses

    monkeypatch.setattr(runtime_module, "get_runtime_statuses", fake_statuses)
    monkeypatch.setattr(inspection_module, "hwp5txt_available", lambda: True)
    monkeypatch.setattr(
        parser_module,
        "extract_hwp_text_with_hwp5txt",
        lambda _path: parser_module.TextExtractionResult(
            succeeded=True,
            text="제1장 총칙\n\n보험계약 안내",
            message="HWP extracted with hwp5txt text route.",
        ),
    )

    with NamedTemporaryFile(suffix=".hwp", delete=False) as handle:
        path = Path(handle.name)
        handle.write(b"hwp-placeholder")

    result = run_pipeline(PipelineRequest(source_path=path, document_format=DocumentFormat.HWP))

    assert result.parser_id == "hwp5txt"
    assert result.decision.value == "degraded_accept"
    assert "degraded_parser_route" in result.handoff.reasons
    assert result.handoff.metadata["parser_route_kind"] == "text_route"


def test_doc_pipeline_uses_antiword_text_fallback_when_selected(monkeypatch) -> None:
    from markbridge.parsers import basic as basic_module

    monkeypatch.setattr(
        basic_module,
        "extract_doc_text_with_antiword",
        lambda _path: basic_module.TextExtractionResult(
            succeeded=True,
            text="1. 안내사항\n\n대리인 접수 가능",
            message="DOC extracted with antiword text fallback.",
        ),
    )

    with NamedTemporaryFile(suffix=".doc", delete=False) as handle:
        path = Path(handle.name)
        handle.write(b"legacy-doc-placeholder")

    result = basic_module.parse_with_current_runtime(
        ParseRequest(source_path=path, document_format=DocumentFormat.DOC),
        "antiword",
    )

    assert result.parser_id == "antiword"
    assert result.document.source_format is DocumentFormat.DOC
    assert any(block.text == "1. 안내사항" for block in result.document.blocks)


def test_hwp_pipeline_uses_hwp5txt_text_route_when_selected(monkeypatch) -> None:
    from markbridge.parsers import basic as basic_module

    monkeypatch.setattr(
        basic_module,
        "extract_hwp_text_with_hwp5txt",
        lambda _path: basic_module.TextExtractionResult(
            succeeded=True,
            text="제1장 총칙\n\n보험계약 안내",
            message="HWP extracted with hwp5txt text route.",
        ),
    )

    with NamedTemporaryFile(suffix=".hwp", delete=False) as handle:
        path = Path(handle.name)
        handle.write(b"hwp-placeholder")

    result = basic_module.parse_with_current_runtime(
        ParseRequest(source_path=path, document_format=DocumentFormat.HWP),
        "hwp5txt",
    )

    assert result.parser_id == "hwp5txt"
    assert result.document.source_format is DocumentFormat.HWP
    assert any(block.text == "제1장 총칙" for block in result.document.blocks)
