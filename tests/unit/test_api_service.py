import json
from dataclasses import replace
from datetime import datetime
from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document as DocxDocument

from markbridge.api.llm import LlmAdvice
from markbridge.api.config import get_settings
from markbridge.api.models import ParseMarkdownExportStatus
from markbridge.api.service import MarkBridgePipeline
from markbridge.pipeline.models import PipelineRequest, PipelineResult
from markbridge.routing.model import LlmUsageMode, RouteLevel, RoutingDecision
from markbridge.shared.ir import DocumentFormat
from markbridge.tracing.model import DisplayExcerpt, IssueSeverity, ParseStatus, ParseTrace, TraceStage
from markbridge.validators.gate import HandoffDecision, QualityGateResult
from markbridge.validators.model import LocationRef, ValidationIssue, ValidationIssueCode, ValidationReport


def test_api_service_upload_smoke() -> None:
    doc = DocxDocument()
    doc.add_heading("API Sample", level=1)
    doc.add_paragraph("Hello from API")
    with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        path = Path(handle.name)
    doc.save(path)

    service = MarkBridgePipeline(get_settings())
    response = service.submit_local_upload(
        filename="sample.docx",
        content=path.read_bytes(),
        llm_requested=False,
    )

    dumped = response.model_dump(mode="json")
    assert dumped["routing"]["primary_parser"] == "python-docx"
    assert dumped["handoff"]["decision"] == "accept"
    assert dumped["trace"]["events"]
    assert any(note.startswith("export_dir=") for note in dumped["notes"])


def test_api_service_hwp_upload_returns_hold() -> None:
    service = MarkBridgePipeline(get_settings())
    response = service.submit_local_upload(
        filename="sample.hwp",
        content=b"hwp-placeholder",
        llm_requested=False,
    )

    dumped = response.model_dump(mode="json")
    assert dumped["source"]["document_format"] == "hwp"
    assert dumped["handoff"]["decision"] == "hold"
    assert dumped["trace"]["events"]


def test_api_service_hwp_upload_can_use_enabled_text_route(monkeypatch) -> None:
    import markbridge.inspection.basic as inspection_module
    import markbridge.parsers.basic as parser_module
    import markbridge.routing.runtime as runtime_module

    original_statuses = runtime_module.get_runtime_statuses

    def fake_statuses():
        statuses = original_statuses()
        statuses["hwp5txt"] = runtime_module.RuntimeParserStatus("hwp5txt", True, True)
        return statuses

    monkeypatch.setattr(runtime_module, "get_runtime_statuses", fake_statuses)
    monkeypatch.setattr(
        inspection_module,
        "hwp5txt_available",
        lambda: True,
    )
    monkeypatch.setattr(
        parser_module,
        "extract_hwp_text_with_hwp5txt",
        lambda _path: parser_module.TextExtractionResult(
            succeeded=True,
            text="제1장 총칙\n\n보험계약 안내",
            message="HWP extracted with hwp5txt text route.",
        ),
    )

    service = MarkBridgePipeline(get_settings())
    response = service.submit_local_upload(
        filename="sample.hwp",
        content=b"hwp-placeholder",
        llm_requested=False,
    )

    dumped = response.model_dump(mode="json")
    assert dumped["source"]["document_format"] == "hwp"
    assert dumped["routing"]["primary_parser"] == "hwp5txt"
    assert dumped["handoff"]["decision"] == "degraded_accept"
    assert "degraded_parser_route" in dumped["handoff"]["reasons"]
    assert dumped["handoff"]["metadata"]["parser_route_kind"] == "text_route"
    assert "parser_route_kind=text_route" in dumped["notes"]
    assert dumped["markdown"]


def test_api_service_build_repair_candidates_adds_line_number() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()

    candidates = service._build_repair_candidates(result)

    assert len(candidates) == 1
    assert candidates[0]["markdown_line_number"] == 12
    assert candidates[0]["patch_proposal"]["markdown_line_number"] == 12


def test_api_service_excerpt_around_focus_prefers_focus_window() -> None:
    service = MarkBridgePipeline(get_settings())

    excerpt = service._excerpt_around_focus(
        "prefix " + ("x" * 120) + " q_{x+t}^{L} " + ("y" * 120) + " suffix",
        focus_text="q_{x+t}^{L}",
        max_chars=80,
    )

    assert "q_{x+t}^{L}" in excerpt
    assert excerpt.startswith("...")
    assert excerpt.endswith("...")


def test_api_service_build_llm_repair_candidates_creates_reviewable_patch() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()
    base_candidates = service._build_repair_candidates(result)
    advice = LlmAdvice(
        used=True,
        raw={
            "repairs": [
                {
                    "issue_id": result.validation.issues[0].issue_id,
                    "candidate_text": "1.3. 해지율( q_{x+t} l )에 관한 사항",
                    "normalized_math": "q_{x+t} l",
                    "confidence": 0.88,
                    "reason": "The surrounding heading indicates actuarial notation for lapse rate.",
                    "uncertain": False,
                }
            ]
        },
    )

    llm_candidates = service._build_llm_repair_candidates(base_candidates, advice=advice)

    assert len(llm_candidates) == 1
    assert llm_candidates[0]["origin"] == "llm"
    assert llm_candidates[0]["strategy"] == "llm_formula_reconstruction"
    assert llm_candidates[0]["patch_proposal"]["replacement_text"] == "1.3. 해지율( q_{x+t} l )에 관한 사항"
    assert llm_candidates[0]["patch_proposal"]["markdown_line_number"] == 12


def test_api_service_maybe_recommend_repair_uses_expanded_output_budget() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()
    repair_candidates = service._build_repair_candidates(result)

    class FakeAdvisor:
        def __init__(self):
            self.max_output_tokens = None

        def recommend_repair(self, *, prompt: str, max_output_tokens: int | None = None):
            self.max_output_tokens = max_output_tokens
            return LlmAdvice(used=True, raw={"repairs": []})

    fake = FakeAdvisor()
    service._llm = fake  # type: ignore[assignment]

    _, record = service._maybe_recommend_repair(
        _build_acquired_source(),
        result,
        repair_candidates=repair_candidates,
        llm_requested=True,
    )

    assert fake.max_output_tokens is not None
    assert fake.max_output_tokens >= 256
    assert record is not None
    assert record["max_output_tokens"] == fake.max_output_tokens


def test_api_service_maybe_recommend_repair_batches_large_target_sets() -> None:
    service = MarkBridgePipeline(get_settings())
    issues = [
        ValidationIssue.create(
            code=ValidationIssueCode.TEXT_CORRUPTION,
            severity=IssueSeverity.WARNING,
            stage=TraceStage.VALIDATION,
            message="Suspicious broken glyphs detected in parsed text.",
            location=LocationRef(block_ref=f"block-{index}", line_hint=f"block {index}"),
            excerpts=(
                DisplayExcerpt(
                    label="broken-text",
                    content=f"broken formula issue-{index}",
                    highlight_text="x",
                    location_hint=f"block {index}",
                ),
            ),
            details={"corruption_class": "formula_placeholder"},
            repairable=True,
        )
        for index in range(9)
    ]
    large_candidates = [
        {
            "issue_id": issue.issue_id,
            "repair_type": "formula_reconstruction",
            "strategy": "llm_required",
            "origin": "deterministic",
            "source_text": f"broken formula {issue.issue_id}",
            "source_span": "x",
            "candidate_text": None,
            "normalized_math": None,
            "confidence": 0.0,
            "rationale": "Need LLM reconstruction.",
            "requires_review": True,
            "llm_recommended": True,
            "block_ref": f"block-{index}",
            "markdown_line_number": index + 1,
            "location_hint": f"block {index}",
            "severity": "warning",
            "patch_proposal": None,
        }
        for index, issue in enumerate(issues)
    ]

    result = replace(
        _build_pipeline_result(),
        validation=ValidationReport(
            issues=tuple(issues),
            summary={"issue_count": len(issues)},
        ),
        metadata={
            **_build_pipeline_result().metadata,
            "markdown": "\n".join(f"line {index}" for index in range(1, 20)),
        },
    )

    class FakeAdvisor:
        def __init__(self):
            self.prompts: list[str] = []

        def recommend_repair(self, *, prompt: str, max_output_tokens: int | None = None):
            self.prompts.append(prompt)
            payload = json.loads(prompt.split("Repair targets:\n", 1)[1])
            repairs = [
                    {
                        "issue_id": item["issue_id"],
                        "candidate_text": f"fixed {item['issue_id']}",
                        "confidence": 0.8,
                    "reason": "batched repair",
                    "uncertain": False,
                }
                for item in payload
            ]
            return LlmAdvice(used=True, raw={"repairs": repairs})

    fake = FakeAdvisor()
    service._llm = fake  # type: ignore[assignment]

    llm_candidates, record = service._maybe_recommend_repair(
        _build_pdf_acquired_source(),
        result,
        repair_candidates=large_candidates,
        llm_requested=True,
    )

    assert len(fake.prompts) == 2
    assert len(llm_candidates) == len(issues)
    assert record is not None
    assert record["batch_count"] == 2
    assert len(record["batches"]) == 2
    assert len(record["generated_candidates"]) == len(issues)


def test_api_service_persists_repair_outputs_with_llm_record(tmp_path: Path) -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result(export_dir=tmp_path)
    base_candidates = service._build_repair_candidates(result)
    llm_record = {
        "created_at": "2026-04-04T00:00:00Z",
        "document_format": "docx",
        "parser_id": "python-docx",
        "targets": [
            {
                "issue_id": result.validation.issues[0].issue_id,
                "source_text": "1.3. 해지율(      )에 관한 사항",
                "source_span": "",
            }
        ],
        "response": {
            "repairs": [
                {
                    "issue_id": result.validation.issues[0].issue_id,
                    "candidate_text": "1.3. 해지율( q_{x+t} l )에 관한 사항",
                }
            ]
        },
        "generated_candidates": [
            {
                "issue_id": result.validation.issues[0].issue_id,
                "origin": "llm",
            }
        ],
    }

    persisted = service._persist_repair_outputs(
        result,
        repair_candidates=base_candidates,
        llm_repair_record=llm_record,
        final_resolved_markdown="# Sample\n\n1.3. 해지율( q x + t l )에 관한 사항",
        final_resolved_patches=[
            {
                "issue_id": result.validation.issues[0].issue_id,
                "action": "replace_text",
                "target_text": "1.3. 해지율(      )에 관한 사항",
                "replacement_text": "1.3. 해지율( q x + t l )에 관한 사항",
            }
        ],
        resolution_summary={
            "repair_issue_count": 1,
            "resolved_issue_count": 1,
            "recovered_deterministic_count": 1,
            "recovered_llm_count": 0,
            "unresolved_repair_issue_count": 0,
            "unresolved_by_class": {},
            "unresolved_by_reason": {},
            "issues": [],
        },
        evaluation={"readiness_score": 72, "readiness_label": "reviewable"},
    )

    repair_path = Path(persisted["repair_candidates_path"])
    llm_path = Path(persisted["llm_formula_repair_path"])
    evaluation_path = Path(persisted["parse_evaluation_path"])
    final_resolved_path = Path(persisted["final_resolved_markdown_path"])
    canonical_path = Path(persisted["canonical_markdown_path"])
    resolution_summary_path = Path(persisted["resolution_summary_path"])
    assert repair_path.exists()
    assert llm_path.exists()
    assert evaluation_path.exists()
    assert final_resolved_path.exists()
    assert canonical_path.exists()
    assert resolution_summary_path.exists()
    repair_body = json.loads(repair_path.read_text(encoding="utf-8"))
    llm_body = json.loads(llm_path.read_text(encoding="utf-8"))
    evaluation_body = json.loads(evaluation_path.read_text(encoding="utf-8"))
    assert repair_body[0]["patch_proposal"]["replacement_text"] == "1.3. 해지율( q x + t l )에 관한 사항"
    assert llm_body["targets"][0]["source_span"] == ""
    assert llm_body["response"]["repairs"][0]["candidate_text"] == "1.3. 해지율( q_{x+t} l )에 관한 사항"
    assert evaluation_body["readiness_label"] == "reviewable"
    assert canonical_path.name == "sample.docx-1.md"
    assert canonical_path.read_text(encoding="utf-8") == "# Sample\n\n1.3. 해지율(      )에 관한 사항"


def test_api_service_build_notes_includes_llm_repair_attempt_metadata() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()

    notes = service._build_notes(
        _build_acquired_source(),
        result,
        LlmAdvice(used=True, recommendation="pypdf"),
        repair_candidates=service._build_repair_candidates(result),
        llm_repair_candidates=[],
        llm_repair_record={
            "targets": [{"issue_id": result.validation.issues[0].issue_id}],
            "generated_candidates": [],
            "error": "empty JSON response",
        },
        persisted_paths={},
    )

    assert "llm_routing_recommendation=pypdf" in notes
    assert "llm_repair_attempted_issues=1" in notes
    assert "llm_repair_generated_candidates=0" in notes
    assert "llm_repair_error=empty JSON response" in notes


def test_api_service_to_response_includes_llm_diagnostics() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()

    response = service._to_response(
        source=_build_acquired_source(),
        result=result,
        llm_requested=True,
        llm_used=True,
        routing_advice=LlmAdvice(used=True, recommendation="pypdf"),
        routing_probe={
            "baseline_parser": "docling",
            "recommended_parser": "pypdf",
            "selected_parser": "docling",
            "override_applied": False,
            "comparison_preview": [
                "baseline=docling score=92.0",
                "recommended=pypdf score=51.0",
                "heading_count 4 -> 0",
                "baseline_retained_after_probe",
            ],
        },
        llm_repair_record={
            "targets": [{"issue_id": result.validation.issues[0].issue_id}],
            "generated_candidates": [],
            "response": {},
            "error": "empty JSON response",
        },
        notes=[],
        repair_candidates=service._build_repair_candidates(result),
        suggested_resolved_markdown=None,
        suggested_resolved_patches=[],
        final_resolved_markdown=None,
        final_resolved_patches=[],
        resolution_summary=service._build_resolution_summary(
            issues=result.validation.issues,
            repair_candidates=service._build_repair_candidates(result),
            final_resolved_patches=[],
            llm_requested=True,
            llm_repair_record={
                "targets": [{"issue_id": result.validation.issues[0].issue_id}],
                "generated_candidates": [],
            },
        ),
        downstream_handoff={
            "policy": "dual_track_review",
            "preferred_markdown_kind": "source",
            "review_required": True,
            "source_markdown_available": True,
            "suggested_resolved_available": False,
            "final_resolved_available": False,
            "rationale": [],
        },
        evaluation={
            "readiness_score": 50,
            "readiness_label": "fragile",
            "issue_count": 1,
            "repair_candidate_count": 1,
            "deterministic_candidate_count": 1,
            "llm_candidate_count": 0,
            "suggested_patch_count": 0,
            "recovered_deterministic_count": 0,
            "recovered_llm_count": 0,
            "unresolved_repair_issue_count": 1,
            "review_required": True,
            "recommended_next_step": "Inspect repair candidates before trusting the parse for downstream indexing.",
            "rationale": [],
        },
        formula_probe_record={
            "page_match": {"page_number": 2},
            "region_image_path": "/tmp/first_formula_probe_region.png",
            "llm_probe": {
                "error": None,
                "response": {
                    "apply_as_patch": False,
                    "confidence": 0.36,
                    "reason": "The crop is still ambiguous.",
                    "replacement_markdown": "candidate preview",
                },
            },
        },
    )

    dumped = response.model_dump(mode="json")
    assert dumped["llm_diagnostics"]["routing_used"] is False
    assert dumped["llm_diagnostics"]["routing_recommendation"] == "pypdf"
    assert dumped["llm_diagnostics"]["routing_baseline_parser"] == "docling"
    assert dumped["llm_diagnostics"]["routing_selected_parser"] == "docling"
    assert dumped["llm_diagnostics"]["routing_override_applied"] is False
    assert dumped["llm_diagnostics"]["routing_comparison_preview"][0] == "baseline=docling score=92.0"
    assert dumped["llm_diagnostics"]["repair_attempted_issues"] == 1
    assert dumped["llm_diagnostics"]["repair_generated_candidates"] == 0
    assert dumped["llm_diagnostics"]["repair_error"] == "empty JSON response"
    assert dumped["llm_diagnostics"]["repair_response_available"] is True
    assert dumped["llm_diagnostics"]["repair_response_preview"] == ["response_present_but_no_repairs"]
    assert dumped["llm_diagnostics"]["formula_probe_attempted"] is True
    assert dumped["llm_diagnostics"]["formula_probe_apply_as_patch"] is False
    assert dumped["llm_diagnostics"]["formula_probe_confidence"] == 0.36
    assert dumped["llm_diagnostics"]["formula_probe_region_image_path"] == "/tmp/first_formula_probe_region.png"
    assert dumped["llm_diagnostics"]["formula_probe_preview"][0] == "matched_page=2"


def test_api_service_build_suggested_resolved_markdown_prefers_llm_patch() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()
    base_candidates = service._build_repair_candidates(result)
    llm_candidates = [
        {
            "issue_id": result.validation.issues[0].issue_id,
            "repair_type": "formula_reconstruction",
            "strategy": "llm_formula_reconstruction",
            "origin": "llm",
            "source_text": "1.3. 해지율(      )에 관한 사항",
            "source_span": "",
            "candidate_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
            "normalized_math": "q_{x+t}^{L}",
            "confidence": 0.74,
            "rationale": "LLM proposed a tighter formula reconstruction.",
            "requires_review": True,
            "llm_recommended": False,
            "block_ref": "block-13",
            "markdown_line_number": 12,
            "location_hint": "block 13",
            "severity": "warning",
            "patch_proposal": {
                "action": "replace_text",
                "target_text": "1.3. 해지율(      )에 관한 사항",
                "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
                "block_ref": "block-13",
                "location_hint": "block 13",
                "markdown_line_number": 12,
                "confidence": 0.74,
                "rationale": "LLM proposed a tighter formula reconstruction.",
                "uncertain": True,
            },
        }
    ]

    resolved_markdown, applied_patches = service._build_suggested_resolved_markdown(
        markdown=str(result.metadata["markdown"]),
        repair_candidates=base_candidates + llm_candidates,
    )

    assert resolved_markdown is not None
    assert "q_{x+t}^L" in resolved_markdown
    assert len(applied_patches) == 1
    assert applied_patches[0]["replacement_text"] == "1.3. 해지율( q_{x+t}^L )에 관한 사항"
    assert applied_patches[0]["issue_id"] == result.validation.issues[0].issue_id


def test_api_service_select_pipeline_result_keeps_baseline_when_probe_score_is_lower(monkeypatch) -> None:
    service = MarkBridgePipeline(get_settings())
    baseline = _build_pipeline_result()
    baseline = replace(
        baseline,
        parser_id="docling",
        metadata={
            **baseline.metadata,
            "markdown": "## Title\n\n## Section\n\nRegular text",
        },
    )
    candidate = replace(
        baseline,
        parser_id="pypdf",
        metadata={
            **baseline.metadata,
            "markdown": "Very long merged line " * 40,
        },
    )

    def fake_run_pipeline_for_parser(source, *, llm_requested, parser_hint, parser_override, llm_route_used):
        return candidate if parser_override == "pypdf" else baseline

    monkeypatch.setattr(service, "_run_pipeline_for_parser", fake_run_pipeline_for_parser)

    selected, probe = service._select_pipeline_result(
        _build_pdf_acquired_source(),
        llm_requested=True,
        parser_hint=None,
        routing_advice=LlmAdvice(used=True, recommendation="pypdf"),
    )

    assert selected.parser_id == "docling"
    assert probe is not None
    assert probe["override_applied"] is False
    assert probe["selected_parser"] == "docling"


def test_api_service_select_pipeline_result_accepts_probe_when_score_is_higher(monkeypatch) -> None:
    service = MarkBridgePipeline(get_settings())
    baseline = _build_pipeline_result()
    baseline = replace(
        baseline,
        parser_id="docling",
        metadata={
            **baseline.metadata,
            "markdown": "Merged line " * 50,
        },
    )
    candidate = replace(
        baseline,
        parser_id="pypdf",
        metadata={
            **baseline.metadata,
            "markdown": "## Title\n\n## Section\n\nRecovered text",
        },
    )

    def fake_run_pipeline_for_parser(source, *, llm_requested, parser_hint, parser_override, llm_route_used):
        return candidate if parser_override == "pypdf" else baseline

    monkeypatch.setattr(service, "_run_pipeline_for_parser", fake_run_pipeline_for_parser)

    selected, probe = service._select_pipeline_result(
        _build_pdf_acquired_source(),
        llm_requested=True,
        parser_hint=None,
        routing_advice=LlmAdvice(used=True, recommendation="pypdf"),
    )

    assert selected.parser_id == "pypdf"
    assert probe is not None
    assert probe["override_applied"] is True
    assert probe["selected_parser"] == "pypdf"


def test_api_service_build_downstream_handoff_prefers_resolved_when_available() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()
    base_candidates = service._build_repair_candidates(result)
    resolution_summary = service._build_resolution_summary(
        issues=result.validation.issues,
        repair_candidates=base_candidates + [
            {
                "issue_id": result.validation.issues[0].issue_id,
                "origin": "llm",
                "patch_proposal": {
                    "target_text": "1.3. 해지율(      )에 관한 사항",
                    "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
                },
            }
        ],
        final_resolved_patches=[
            {
                "issue_id": result.validation.issues[0].issue_id,
                "origin": "llm",
                "action": "replace_text",
                "target_text": "1.3. 해지율(      )에 관한 사항",
                "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
            }
        ],
        llm_requested=True,
        llm_repair_record={
            "targets": [{"issue_id": result.validation.issues[0].issue_id}],
            "generated_candidates": [],
        },
    )
    handoff = service._build_downstream_handoff(
        handoff_decision=result.decision.value,
        markdown=str(result.metadata["markdown"]),
        final_resolved_markdown="# Sample\n\n1.3. 해지율( q_{x+t}^L )에 관한 사항",
        final_resolved_patches=[
            {
                "issue_id": result.validation.issues[0].issue_id,
                "origin": "llm",
                "action": "replace_text",
                "target_text": "1.3. 해지율(      )에 관한 사항",
                "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
                "uncertain": False,
            }
        ],
        repair_candidates=base_candidates + [
            {
                "issue_id": result.validation.issues[0].issue_id,
                "origin": "llm",
                "patch_proposal": {
                    "target_text": "1.3. 해지율(      )에 관한 사항",
                    "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
                    "uncertain": False,
                },
            }
        ],
        resolution_summary=resolution_summary,
    )

    assert handoff["policy"] == "resolved_preferred"
    assert handoff["preferred_markdown_kind"] == "resolved"
    assert handoff["review_required"] is False
    assert handoff["suggested_resolved_available"] is True
    assert handoff["final_resolved_available"] is True


def test_api_service_build_downstream_handoff_keeps_source_canonical_when_placeholder_remains() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()
    base_candidates = service._build_repair_candidates(result)
    resolution_summary = service._build_resolution_summary(
        issues=result.validation.issues,
        repair_candidates=base_candidates,
        final_resolved_patches=[
            {
                "issue_id": result.validation.issues[0].issue_id,
                "origin": "deterministic",
                "action": "replace_text",
                "target_text": "1.3. 해지율(      )에 관한 사항",
                "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
                "uncertain": False,
            }
        ],
        llm_requested=False,
    )

    handoff = service._build_downstream_handoff(
        handoff_decision=result.decision.value,
        markdown=str(result.metadata["markdown"]),
        final_resolved_markdown="# Sample\n\n1.3. 해지율( q_{x+t}^L )에 관한 사항\n\n<!-- formula-not-decoded -->",
        final_resolved_patches=[
            {
                "issue_id": result.validation.issues[0].issue_id,
                "origin": "deterministic",
                "action": "replace_text",
                "target_text": "1.3. 해지율(      )에 관한 사항",
                "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
                "uncertain": False,
            }
        ],
        repair_candidates=base_candidates,
        resolution_summary=resolution_summary,
    )

    assert handoff["policy"] == "dual_track_review"
    assert handoff["preferred_markdown_kind"] == "source"
    assert handoff["review_required"] is True
    assert handoff["final_resolved_available"] is True
    assert handoff["remaining_placeholder_count"] == 1


def test_api_service_build_suggested_resolved_markdown_falls_back_to_next_applicable_candidate() -> None:
    service = MarkBridgePipeline(get_settings())

    resolved_markdown, applied_patches = service._build_suggested_resolved_markdown(
        markdown="# Sample\n\n1.3. 해지율(      )에 관한 사항",
        repair_candidates=[
            {
                "issue_id": "issue-1",
                "origin": "llm",
                "confidence": 0.95,
                "patch_proposal": {
                    "action": "replace_text",
                    "target_text": "missing target",
                    "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
                    "markdown_line_number": 3,
                },
            },
            {
                "issue_id": "issue-1",
                "origin": "deterministic",
                "confidence": 0.7,
                "patch_proposal": {
                    "action": "replace_text",
                    "target_text": "1.3. 해지율(      )에 관한 사항",
                    "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
                    "markdown_line_number": 3,
                },
            },
        ],
    )

    assert resolved_markdown is not None
    assert "q_{x+t}^L" in resolved_markdown
    assert applied_patches[0]["origin"] == "deterministic"


def test_api_service_build_parse_evaluation_marks_reviewable_state() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()
    resolution_summary = service._build_resolution_summary(
        issues=result.validation.issues,
        repair_candidates=service._build_repair_candidates(result),
        final_resolved_patches=[],
        llm_requested=False,
    )
    evaluation = service._build_parse_evaluation(
        issue_count=len(result.validation.issues),
        repair_candidates=service._build_repair_candidates(result),
        final_resolved_patches=[],
        downstream_handoff={
            "policy": "dual_track_review",
            "preferred_markdown_kind": "source",
            "review_required": True,
            "source_markdown_available": True,
            "suggested_resolved_available": False,
            "final_resolved_available": False,
            "rationale": [],
        },
        resolution_summary=resolution_summary,
    )

    assert evaluation["readiness_label"] == "reviewable"
    assert evaluation["repair_candidate_count"] == 1
    assert evaluation["deterministic_candidate_count"] == 1
    assert evaluation["unresolved_repair_issue_count"] == 1
    assert any("Deterministic candidates still recommending LLM review" in item for item in evaluation["rationale"])


def test_api_service_build_parse_evaluation_penalizes_placeholder_residue() -> None:
    service = MarkBridgePipeline(get_settings())
    evaluation = service._build_parse_evaluation(
        issue_count=1,
        repair_candidates=[
            {
                "issue_id": "issue-1",
                "origin": "llm",
                "requires_review": True,
                "llm_recommended": False,
            }
        ],
        final_resolved_patches=[
            {
                "issue_id": "issue-1",
                "origin": "llm",
                "action": "replace_text",
                "target_text": "broken",
                "replacement_text": "fixed",
            }
        ],
        downstream_handoff={
            "policy": "dual_track_review",
            "preferred_markdown_kind": "source",
            "review_required": True,
            "source_markdown_available": True,
            "suggested_resolved_available": True,
            "final_resolved_available": True,
            "remaining_placeholder_count": 3,
            "rationale": [],
        },
        resolution_summary={
            "recovered_llm_count": 1,
            "unresolved_repair_issue_count": 0,
        },
    )

    assert evaluation["readiness_label"] == "reviewable"
    assert any("Formula placeholders still remain in final resolved markdown: 3" in item for item in evaluation["rationale"])


def test_api_service_build_parse_evaluation_counts_strong_deterministic_repairs() -> None:
    service = MarkBridgePipeline(get_settings())
    resolution_summary = service._build_resolution_summary(
        issues=(),
        repair_candidates=[
            {
                "issue_id": "issue-1",
                "origin": "deterministic",
                "requires_review": True,
                "llm_recommended": False,
                "patch_proposal": {
                    "target_text": "  ",
                    "replacement_text": "q_{x+t}^{L}",
                },
            }
        ],
        final_resolved_patches=[
            {
                "issue_id": "issue-1",
                "origin": "deterministic",
                "action": "replace_text",
                "target_text": "  ",
                "replacement_text": "q_{x+t}^{L}",
            }
        ],
        llm_requested=False,
    )
    evaluation = service._build_parse_evaluation(
        issue_count=1,
        repair_candidates=[
            {
                "issue_id": "issue-1",
                "origin": "deterministic",
                "requires_review": True,
                "llm_recommended": False,
            }
        ],
        final_resolved_patches=[
            {
                "issue_id": "issue-1",
                "origin": "deterministic",
                "action": "replace_text",
                "target_text": "  ",
                "replacement_text": "q_{x+t}^{L}",
            }
        ],
        downstream_handoff={
            "policy": "resolved_preferred",
            "preferred_markdown_kind": "resolved",
            "review_required": False,
            "source_markdown_available": True,
            "suggested_resolved_available": True,
            "final_resolved_available": True,
            "rationale": [],
        },
        resolution_summary=resolution_summary,
    )

    assert evaluation["readiness_label"] == "ready"
    assert evaluation["recovered_deterministic_count"] == 1
    assert any("High-structure deterministic candidates" in item for item in evaluation["rationale"])


def test_api_service_build_resolution_summary_tracks_reason_and_class() -> None:
    service = MarkBridgePipeline(get_settings())
    issue = ValidationIssue.create(
        code=ValidationIssueCode.TEXT_CORRUPTION,
        severity=IssueSeverity.WARNING,
        stage=TraceStage.VALIDATION,
        message="Formula placeholder remained unresolved.",
        location=LocationRef(block_ref="block-77", line_hint="block 77"),
        excerpts=(
            DisplayExcerpt(
                label="broken-text",
                content="보험료 산식 [formula:premium_01]",
                highlight_text="[formula:premium_01]",
                location_hint="block 77",
            ),
        ),
        details={"corruption_class": "formula_placeholder"},
        repairable=True,
    )

    resolution_summary = service._build_resolution_summary(
        issues=(issue,),
        repair_candidates=[
            {
                "issue_id": issue.issue_id,
                "origin": "deterministic",
                "strategy": "llm_required",
                "confidence": 0.0,
                "llm_recommended": True,
                "patch_proposal": None,
            }
        ],
        final_resolved_patches=[],
        llm_requested=False,
    )

    assert resolution_summary["unresolved_repair_issue_count"] == 1
    assert resolution_summary["unresolved_by_class"] == {"formula_placeholder": 1}
    assert resolution_summary["unresolved_by_reason"] == {"llm_not_requested": 1}
    assert resolution_summary["issues"][0]["unresolved_reason"] == "llm_not_requested"


def test_api_service_build_resolution_summary_tracks_selection_reason_and_candidate_decisions() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()
    deterministic_candidates = service._build_repair_candidates(result)
    llm_candidates = [
        {
            "issue_id": result.validation.issues[0].issue_id,
            "repair_type": "formula_reconstruction",
            "strategy": "llm_formula_reconstruction",
            "origin": "llm",
            "source_text": "1.3. 해지율(      )에 관한 사항",
            "candidate_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
            "confidence": 0.91,
            "rationale": "LLM proposed a tighter formula reconstruction.",
            "patch_proposal": {
                "action": "replace_text",
                "target_text": "1.3. 해지율(      )에 관한 사항",
                "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
            },
        }
    ]

    resolution_summary = service._build_resolution_summary(
        issues=result.validation.issues,
        repair_candidates=deterministic_candidates + llm_candidates,
        final_resolved_patches=[
            {
                "issue_id": result.validation.issues[0].issue_id,
                "origin": "llm",
                "action": "replace_text",
                "target_text": "1.3. 해지율(      )에 관한 사항",
                "replacement_text": "1.3. 해지율( q_{x+t}^L )에 관한 사항",
            }
        ],
        llm_requested=True,
        llm_repair_record={
            "targets": [{"issue_id": result.validation.issues[0].issue_id}],
            "generated_candidates": llm_candidates,
        },
    )

    issue_summary = resolution_summary["issues"][0]

    assert issue_summary["selected_origin"] == "llm"
    assert issue_summary["selection_reason"] == "llm_priority"
    assert issue_summary["selected_confidence"] == 0.91
    assert issue_summary["candidate_decisions"][0]["selected"] is False
    assert issue_summary["candidate_decisions"][0]["rejected_reason"] == "lower_priority_origin"
    assert issue_summary["candidate_decisions"][1]["selected"] is True


def test_api_service_build_resolution_summary_marks_llm_attempted_without_generated_candidate() -> None:
    service = MarkBridgePipeline(get_settings())
    issue = ValidationIssue.create(
        code=ValidationIssueCode.TEXT_CORRUPTION,
        severity=IssueSeverity.WARNING,
        stage=TraceStage.VALIDATION,
        message="Formula placeholder remained unresolved after LLM attempt.",
        location=LocationRef(block_ref="block-88", line_hint="block 88"),
        excerpts=(
            DisplayExcerpt(
                label="broken-text",
                content="보험료 산식 [formula:reserve_02]",
                highlight_text="[formula:reserve_02]",
                location_hint="block 88",
            ),
        ),
        details={"corruption_class": "formula_placeholder"},
        repairable=True,
    )

    resolution_summary = service._build_resolution_summary(
        issues=(issue,),
        repair_candidates=[
            {
                "issue_id": issue.issue_id,
                "origin": "deterministic",
                "strategy": "llm_required",
                "confidence": 0.0,
                "llm_recommended": True,
                "patch_proposal": None,
            }
        ],
        final_resolved_patches=[],
        llm_requested=True,
        llm_repair_record={
            "targets": [{"issue_id": issue.issue_id}],
            "generated_candidates": [],
        },
    )

    assert resolution_summary["unresolved_by_reason"] == {"llm_no_repair_generated": 1}
    assert resolution_summary["issues"][0]["llm_attempted"] is True


def test_api_service_build_resolution_summary_marks_selected_patch_not_applied() -> None:
    service = MarkBridgePipeline(get_settings())
    result = _build_pipeline_result()

    resolution_summary = service._build_resolution_summary(
        issues=result.validation.issues,
        repair_candidates=service._build_repair_candidates(result),
        final_resolved_patches=[],
        llm_requested=False,
        llm_repair_record=None,
    )

    assert resolution_summary["issues"][0]["selection_reason"] == "only_patch_proposal"
    assert resolution_summary["issues"][0]["selected_origin"] == "deterministic"
    assert resolution_summary["issues"][0]["candidate_decisions"][0]["selected"] is True
    assert resolution_summary["unresolved_by_reason"] == {"selected_patch_not_applied": 1}


def test_api_service_build_resolution_summary_marks_best_applicable_patch_when_lower_rank_wins() -> None:
    service = MarkBridgePipeline(get_settings())

    resolution_summary = service._build_resolution_summary(
        issues=(),
        repair_candidates=[
            {
                "issue_id": "issue-1",
                "origin": "llm",
                "strategy": "llm_formula_reconstruction",
                "confidence": 0.95,
                "patch_proposal": {
                    "target_text": "missing target",
                    "replacement_text": "q_{x+t}^L",
                },
            },
            {
                "issue_id": "issue-1",
                "origin": "deterministic",
                "strategy": "deterministic_transliteration_with_llm_review",
                "confidence": 0.7,
                "patch_proposal": {
                    "target_text": "  ",
                    "replacement_text": "q_{x+t}^L",
                },
            },
        ],
        final_resolved_patches=[
            {
                "issue_id": "issue-1",
                "origin": "deterministic",
                "target_text": "  ",
                "replacement_text": "q_{x+t}^L",
            }
        ],
        llm_requested=True,
    )

    issue_summary = resolution_summary["issues"][0]

    assert issue_summary["selected_origin"] == "deterministic"
    assert issue_summary["selection_reason"] == "best_applicable_patch"
    assert issue_summary["candidate_decisions"][0]["rejected_reason"] == "patch_not_applicable"
    assert issue_summary["candidate_decisions"][1]["selected"] is True


def test_api_service_lists_and_reads_parse_markdown_exports(tmp_path: Path, monkeypatch) -> None:
    _write_export_run(
        tmp_path=tmp_path,
        run_id="run-1",
        source_name="300233_계약관계자변경.docx",
        status="succeeded",
        created_at="2026-04-15T02:53:39Z",
        canonical_markdown=(
            "## 1) 계약자변경 유의사항 및 공통사항\n"
            "> 계약자 변경 시 유의사항\n\n"
            "## 3) 계약자 변경 구비서류\n"
            "| 접수방법 |  | 구비서류 안내 |\n"
            "| --- | --- | --- |\n"
            "| 우편/ FC방문 접수 |  | ... |"
        ),
        line_map=[
            {"line_number": 1, "text": "## 1) 계약자변경 유의사항 및 공통사항", "refs": ["block-1"], "page_number": 1},
            {"line_number": 2, "text": "> 계약자 변경 시 유의사항", "refs": ["block-2"], "page_number": 1},
            {"line_number": 4, "text": "## 3) 계약자 변경 구비서류", "refs": ["block-3"], "page_number": 1},
            {"line_number": 5, "text": "| 접수방법 |  | 구비서류 안내 |", "refs": ["block-4"], "page_number": 1},
            {"line_number": 6, "text": "| --- | --- | --- |", "refs": ["block-4"], "page_number": 1},
            {"line_number": 7, "text": "| 우편/ FC방문 접수 |  | ... |", "refs": ["block-4"], "page_number": 1},
        ],
    )
    _write_export_run(
        tmp_path=tmp_path,
        run_id="run-2",
        source_name="300138_라이프앱가능업무.docx",
        status="failed",
        created_at="2026-04-15T02:10:11Z",
        canonical_markdown="# failed",
        line_map=[],
    )
    _write_export_run(
        tmp_path=tmp_path,
        run_id="run-3",
        source_name="sample.docx",
        status="succeeded",
        created_at="2026-04-15T02:59:11Z",
        canonical_markdown="# sample",
        line_map=[],
    )
    _write_export_run(
        tmp_path=tmp_path,
        run_id="run-4",
        source_name="markbridge_0sudnbbw.docx",
        status="succeeded",
        created_at="2026-04-15T03:01:00Z",
        canonical_markdown="# temp",
        line_map=[],
    )
    monkeypatch.setenv("MARKBRIDGE_WORK_DIR", str(tmp_path))

    service = MarkBridgePipeline(get_settings())

    listing = service.list_parse_markdown_exports(
        updated_after=datetime.fromisoformat("2026-04-15T00:00:00+00:00"),
        parse_status=ParseMarkdownExportStatus.COMPLETED,
    )

    assert len(listing.items) == 1
    item = listing.items[0]
    assert item.document_name == "300233_계약관계자변경.docx"
    assert item.canonical_markdown_name == "300233_계약관계자변경.docx-1.md"
    assert item.parse_status.value == "completed"
    assert item.document_id.startswith("doc_")

    document, content, etag = service.get_parse_markdown_content(item.document_id)
    assert document.document_name == item.document_name
    assert "## 3) 계약자 변경 구비서류" in content
    assert len(etag) == 64

    blocks = service.list_parse_markdown_blocks(item.document_id)
    assert [block.block_kind for block in blocks.blocks] == ["heading", "note", "heading", "table"]

    block_document, block_content, block_etag = service.get_parse_markdown_block_content(
        item.document_id,
        "block-0004",
    )
    assert block_document.document_id == item.document_id
    assert block_content.startswith("| 접수방법")
    assert len(block_etag) == 64


def _build_pipeline_result(export_dir: Path | None = None) -> PipelineResult:
    issue = ValidationIssue.create(
        code=ValidationIssueCode.TEXT_CORRUPTION,
        severity=IssueSeverity.WARNING,
        stage=TraceStage.VALIDATION,
        message="Suspicious broken glyphs detected in parsed text.",
        location=LocationRef(block_ref="block-13", line_hint="block 13"),
        excerpts=(
            DisplayExcerpt(
                label="broken-text",
                content="1.3. 해지율(      )에 관한 사항",
                highlight_text="",
                location_hint="block 13",
            ),
        ),
        details={"corruption_class": "inline_formula_corruption"},
        repairable=True,
    )
    request = PipelineRequest(
        source_path=Path("/tmp/sample.docx"),
        document_format=DocumentFormat.DOCX,
    )
    trace = ParseTrace.create(source_path=request.source_path, document_format=request.document_format)
    return PipelineResult.create(
        request=request,
        trace=trace,
        route=RoutingDecision(
            level=RouteLevel.DETERMINISTIC_ONLY,
            primary_parser="python-docx",
            llm_usage=LlmUsageMode.NONE,
        ),
        validation=ValidationReport(issues=(issue,), summary={"issue_count": 1}),
        handoff=QualityGateResult(
            decision=HandoffDecision.DEGRADED_ACCEPT,
            summary="Allow downstream handoff with degraded status because warning-level validation issues exist.",
        ),
        parser_id="python-docx",
        status=ParseStatus.DEGRADED,
        metadata={
            "markdown": "# Sample\n\n1.3. 해지율(      )에 관한 사항",
            "markdown_line_map": [
                {"line_number": 1, "text": "# Sample", "refs": ["block-1"]},
                {"line_number": 12, "text": "1.3. 해지율(      )에 관한 사항", "refs": ["block-13", "block 13"]},
            ],
            "repair_candidates": [
                {
                    "issue_id": issue.issue_id,
                    "repair_type": "formula_reconstruction",
                    "strategy": "deterministic_transliteration_with_llm_review",
                    "origin": "deterministic",
                    "source_text": "1.3. 해지율(      )에 관한 사항",
                    "source_span": "",
                    "candidate_text": "1.3. 해지율( q x + t l )에 관한 사항",
                    "normalized_math": "q_{x+t} l",
                    "confidence": 0.5,
                    "rationale": "Formula-like corruption requires review.",
                    "requires_review": True,
                    "llm_recommended": True,
                    "block_ref": "block-13",
                    "location_hint": "block 13",
                    "severity": "warning",
                    "patch_proposal": {
                        "action": "replace_text",
                        "target_text": "1.3. 해지율(      )에 관한 사항",
                        "replacement_text": "1.3. 해지율( q x + t l )에 관한 사항",
                        "block_ref": "block-13",
                        "location_hint": "block 13",
                        "confidence": 0.5,
                        "rationale": "Formula-like corruption requires review.",
                        "uncertain": True,
                    },
                }
            ],
            "export_dir": str(export_dir) if export_dir is not None else "/tmp/markbridge-run",
            "source_name": "sample.docx",
        },
    )


def _build_acquired_source():
    from markbridge.api.models import SourceKind
    from markbridge.api.service import AcquiredSource

    return AcquiredSource(
        source_kind=SourceKind.UPLOAD,
        source_name="sample.docx",
        uri=None,
        path=Path("/tmp/sample.docx"),
        document_format=DocumentFormat.DOCX,
        size_bytes=128,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _build_pdf_acquired_source():
    from markbridge.api.models import SourceKind
    from markbridge.api.service import AcquiredSource

    return AcquiredSource(
        source_kind=SourceKind.S3_URI,
        source_name="sample.pdf",
        uri="s3://bucket/sample.pdf",
        path=Path("/tmp/sample.pdf"),
        document_format=DocumentFormat.PDF,
        size_bytes=256,
        content_type="application/pdf",
    )


def _write_export_run(
    *,
    tmp_path: Path,
    run_id: str,
    source_name: str,
    status: str,
    created_at: str,
    canonical_markdown: str,
    line_map: list[dict[str, object]],
) -> None:
    run_dir = tmp_path / run_id
    run_dir.mkdir(parents=True, exist_ok=True)
    manifest = {
        "run_id": run_id,
        "run_dir": str(run_dir),
        "created_at": created_at,
        "metadata": {
            "source_name": source_name,
            "status": status,
        },
    }
    (run_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
    (run_dir / f"{source_name}-1.md").write_text(canonical_markdown, encoding="utf-8")
    (run_dir / "markdown_line_map.json").write_text(json.dumps(line_map, ensure_ascii=False), encoding="utf-8")
