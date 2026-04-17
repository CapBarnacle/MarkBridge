"""Basic deterministic inspectors for currently enabled formats."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from markbridge.inspection.model import DocConversionFeatures, HwpInspectionFeatures
from markbridge.inspection.model import (
    CommonInspectionFeatures,
    DocxInspectionFeatures,
    InspectionReport,
    PdfInspectionFeatures,
    XlsxInspectionFeatures,
)
from markbridge.parsers.conversion import antiword_available, hwp5txt_available, libreoffice_available
from markbridge.shared.ir import DocumentFormat


def inspect_document(source_path: Path, document_format: DocumentFormat) -> InspectionReport:
    source_path = Path(source_path)
    common = CommonInspectionFeatures(
        file_size_bytes=source_path.stat().st_size if source_path.exists() else None,
    )
    if document_format is DocumentFormat.PDF:
        reader = PdfReader(str(source_path))
        page_count = len(reader.pages)
        text_pages = 0
        table_candidates = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                text_pages += 1
            table_candidates += text.count("|")
        return InspectionReport(
            source_path=source_path,
            document_format=document_format,
            common=CommonInspectionFeatures(
                file_size_bytes=common.file_size_bytes,
                page_count=page_count,
                complexity_score=float(table_candidates > 0),
            ),
            pdf=PdfInspectionFeatures(
                text_layer_coverage=(text_pages / page_count) if page_count else None,
                table_candidate_count=table_candidates or None,
            ),
        )

    if document_format is DocumentFormat.DOCX:
        doc = DocxDocument(str(source_path))
        paragraphs = list(doc.paragraphs)
        table_count = len(doc.tables)
        heading_styles = sum(1 for p in paragraphs if (p.style and "Heading" in p.style.name))
        return InspectionReport(
            source_path=source_path,
            document_format=document_format,
            common=CommonInspectionFeatures(
                file_size_bytes=common.file_size_bytes,
                complexity_score=float(table_count > 0),
            ),
            docx=DocxInspectionFeatures(
                heading_style_availability=heading_styles > 0,
                paragraph_count=len(paragraphs),
                table_count=table_count,
            ),
        )

    if document_format is DocumentFormat.XLSX:
        workbook = load_workbook(str(source_path), data_only=False)
        sheets = workbook.worksheets
        merged_count = sum(len(sheet.merged_cells.ranges) for sheet in sheets)
        formula_cells = 0
        non_empty_cells = 0
        for sheet in sheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        non_empty_cells += 1
                        if isinstance(cell.value, str) and cell.value.startswith("="):
                            formula_cells += 1
        return InspectionReport(
            source_path=source_path,
            document_format=document_format,
            common=CommonInspectionFeatures(
                file_size_bytes=common.file_size_bytes,
                sheet_count=len(sheets),
                complexity_score=float(merged_count > 0 or formula_cells > 0),
            ),
            xlsx=XlsxInspectionFeatures(
                sheet_count=len(sheets),
                merged_cell_count=merged_count,
                formula_ratio=(formula_cells / non_empty_cells) if non_empty_cells else None,
            ),
        )

    if document_format is DocumentFormat.DOC:
        libreoffice_ready = libreoffice_available()
        antiword_ready = antiword_available()
        quality_signals: list[str] = []
        warnings: list[str] = []
        if libreoffice_ready:
            quality_signals.append("libreoffice_route_available")
        if antiword_ready:
            quality_signals.append("antiword_text_route_available")
        if not quality_signals:
            quality_signals.append("conversion_tool_missing")
            warnings.append("No DOC execution route is available in the current runtime.")
        return InspectionReport(
            source_path=source_path,
            document_format=document_format,
            common=common,
            doc=DocConversionFeatures(
                conversion_feasibility=libreoffice_ready or antiword_ready,
                conversion_output_quality_signals=tuple(quality_signals),
            ),
            warnings=tuple(warnings),
        )

    if document_format is DocumentFormat.HWP:
        hwp_ready = hwp5txt_available()
        return InspectionReport(
            source_path=source_path,
            document_format=document_format,
            common=common,
            hwp=HwpInspectionFeatures(
                execution_feasibility=hwp_ready,
                execution_route_candidates=("hwp5txt",) if hwp_ready else (),
            ),
            warnings=() if hwp_ready else ("HWP parsing is not implemented in the current runtime.",),
        )

    return InspectionReport(
        source_path=source_path,
        document_format=document_format,
        common=common,
        warnings=("inspection not implemented for this format",),
    )
